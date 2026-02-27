"""
Utilities for reading and writing DOCX resumes while preserving layout.

We use python-docx to:
- Extract plain text for the LLM pipeline
- Track which paragraphs are bullets so we can rewrite them in-place
"""

from pathlib import Path
from typing import Dict, List, Tuple, Union

from docx import Document  # type: ignore[import-untyped]


PathLike = Union[str, Path]


def read_docx_to_text_and_map(path: PathLike) -> Tuple[str, Dict[str, List[int]]]:
    """
    Read a DOCX resume and return:
    - plain text (for LLM parsing)
    - a map: original_bullet_text -> list of paragraph indices that contain it
    """
    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)

    lines: List[str] = []
    bullet_map: Dict[str, List[int]] = {}

    for idx, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue

        lines.append(text)

        # Map every non-empty paragraph text to its indices.
        # This is more robust across different Word styles where bullets
        # may not use a specific "bullet" style name.
        bullet_map.setdefault(text, []).append(idx)

    full_text = "\n".join(lines)
    return full_text, bullet_map


def write_optimized_docx(
    original_path: PathLike,
    output_path: PathLike,
    bullet_map: Dict[str, List[int]],
    rewritten_bullets: List[dict],
) -> None:
    """
    Open original DOCX and replace bullet paragraphs whose text matches
    rewriter.original with rewriter.rewritten.

    This keeps the overall document layout (sections, fonts, etc.) while
    updating only the bullet text.
    """
    doc = Document(str(original_path))
    paragraphs = list(doc.paragraphs)

    replace_map = {
        b["original"]: b["rewritten"] for b in rewritten_bullets
    }

    for original, rewritten in replace_map.items():
        indices = bullet_map.get(original, [])
        for idx in indices:
            if 0 <= idx < len(paragraphs):
                paragraphs[idx].text = rewritten

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

