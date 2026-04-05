"""Export optimized resume to DOCX."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph(doc, text, bold=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(14 if level == 1 else 12)
    return h

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Name - large and bold
    p = doc.add_paragraph()
    r = p.add_run("SUSITH HEMATHILAKA")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Calibri"

    # Title line
    add_paragraph(doc, "Senior Full Stack Developer | Java | React | AWS | Microservices", font_size=11)
    add_paragraph(doc, "+971-562-877-944 | Al Karama, Dubai | susith.rsj@gmail.com", font_size=10)
    add_paragraph(doc, "linkedin.com/in/susithrj | github.com/susithrj | medium.com/susithrj | susithrj.com", font_size=10)
    add_paragraph(doc, "Visa Status: Available to start immediately", bold=True, font_size=10)
    doc.add_paragraph()

    add_heading(doc, "PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(
        "Senior Full Stack Developer with 6+ years of experience building scalable, secure full-stack solutions using Java, React, and AWS, with expertise in data orchestration and integration and a strong background in open finance and fintech. Proven track record designing cloud-native systems, mentoring engineers, and driving technical leadership for mission-critical platforms (Fortune 500, Saudi financial institutions). Strong focus on production stability, code quality, and Agile delivery."
    )
    doc.add_paragraph()

    add_heading(doc, "KEY ACHIEVEMENTS", level=1)
    achievements = [
        "Led 100% of backend components for critical production go-lives serving major Saudi financial clients",
        "Monitored system performance and debugged 250+ complex production issues using metrics, logging, and tracing, ensuring zero data loss",
        "Achieved 100% throughput improvement through distributed system design with JGroups failover mechanisms",
        "Reduced operational overhead by 2x through cost-optimized ETL and data sync solutions on AWS Serverless (data orchestration)",
        "Optimized bulk payment processing by 20x using PL/SQL stored procedures for high-volume transactions",
    ]
    for a in achievements:
        doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph()

    add_heading(doc, "TECHNICAL SKILLS", level=1)
    skills = [
        "Backend: Java 8/17/21, Spring Boot, Spring Framework, Hibernate, Spring Security, Microservices, REST APIs, Event-Driven Architecture, Node.js, Python",
        "Frontend: React, Reactjs, Angular, Vue.js, TypeScript, JavaScript (ES6+), HTML5, CSS3",
        "Cloud & Platform: AWS (EC2, ECS, Lambda, S3, RDS, CloudFormation, Step Functions), Azure, scalable backend services",
        "Database: MongoDB, Oracle, MySQL, PostgreSQL, Redis, NoSQL, PL/SQL, schema design, large dataset handling",
        "Architecture: Microservices, RESTful services, data orchestration and integration, distributed systems, fault-tolerant design, Kafka, RabbitMQ",
        "CI/CD & Monitoring: Jenkins, GitHub Actions, Datadog, CloudWatch, ELK, SonarQube, JMeter",
        "Domains: Open finance, FinTech, enterprise platforms, mission-critical systems",
    ]
    for s in skills:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph()

    add_heading(doc, "PROFESSIONAL EXPERIENCE", level=1)

    # Sysco LABS
    p = doc.add_paragraph()
    p.add_run("Associate Technical Lead").bold = True
    p.add_run(" | February 2025 – December 2025")
    p = doc.add_paragraph()
    p.add_run("Sysco LABS,").bold = True
    p.add_run(" Sri Lanka")
    sysco_bullets = [
        "Designed and implemented a cloud-native, scalable, and secure full-stack solution for a Fortune 500 enterprise, utilizing Java, React, and AWS to modernize the pricing platform and support nationwide clients.",
        "Architected and delivered high-impact, cloud-native features for 40,000+ enterprise users, leveraging Java, React, and AWS Serverless to define technical roadmaps and drive business growth.",
        "Orchestrated end-to-end development lifecycle from system architecture through production deployment using AWS serverless (Lambda, Step Functions), establishing secure, resilient, and scalable applications and data orchestration and integration across services.",
        "Drove strategic POC development for cloud-native data sync workflows and ETL pipelines using AWS Glue, implementing cost-optimized solutions that reduced operational overhead by 2x and demonstrated data orchestration and integration expertise.",
        "Mentored engineers through architectural guidance, code reviews, and knowledge sharing, focusing on cloud-native and full-stack best practices, improving code quality and technical leadership.",
    ]
    for b in sysco_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # Intervest
    p = doc.add_paragraph()
    p.add_run("Senior Software Engineer").bold = True
    p.add_run(" | April 2023 – February 2025")
    p = doc.add_paragraph()
    p.add_run("Intervest,").bold = True
    p.add_run(" Sri Lanka")
    intervest_bullets = [
        "Contributed to revamping travel insurance platform, migrating from monolith to microservices using Strangler Fig and BFF patterns.",
        "Implemented distributed caching with Redis and deployed microservices on AWS ECS to improve scalability and response times.",
        "Developed Address Service integration and contributed to iterative client rollouts, supporting platform expansion across multiple enterprise clients.",
    ]
    for b in intervest_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # DirectFN (Senior)
    p = doc.add_paragraph()
    p.add_run("DirectFN,").bold = True
    p.add_run(" Sri Lanka")
    directfn_sr_bullets = [
        "Led production go-live projects for Musharaka Capital and post-go-live operations for Bank Saudi Fransi, owning 100% of backend components (Java services).",
        "Troubleshot 250+ critical production incidents across Java backends (connection pool exhaustion, cache sync, transaction deadlocks).",
        "Built FIX-to-FIXML conversion for settlement management, enabling seamless communication with post-trade clearing parties (Edda/Muqassa).",
        "Collaborated with 30+ stakeholders and business relationship managers for Bank Saudi Fransi and Yaqeen Capital on compliance and feature delivery.",
        "Mentored three software engineers through code reviews, pair programming, and domain knowledge transfer.",
    ]
    for b in directfn_sr_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # Software Engineer - DirectFN
    p = doc.add_paragraph()
    p.add_run("Software Engineer").bold = True
    p.add_run(" | June 2021 – April 2023")
    p = doc.add_paragraph()
    p.add_run("DirectFN,").bold = True
    p.add_run(" Sri Lanka")
    directfn_bullets = [
        "Designed a distributed trading system with failover mechanisms that improved throughput by 100%, supporting 2x growth in customer onboarding.",
        "Optimized bulk payments using PL/SQL stored procedures with set-based operations (bulk collect, FORALL), reducing processing times by 20x.",
        "Provided on-call production support for six Saudi clients with collective daily turnover exceeding 5 billion SAR, ensuring 24/7 availability.",
        "Developed GTN (Global Trading Network) integration to boost client access to international markets, resulting in 25% increase in trading cash flow.",
        "Strengthened code quality and reliability with comprehensive unit and integration test suites for core trading services.",
        "Authored changelogs, release notes, and upgrade guides for smooth deployment and adoption of new features.",
    ]
    for b in directfn_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # Ones&Zeros
    p = doc.add_paragraph()
    p.add_run("Associate Software Engineer").bold = True
    p.add_run(" | July 2020 – June 2021")
    p = doc.add_paragraph()
    p.add_run("Ones&Zeros,").bold = True
    p.add_run(" Sri Lanka")
    oz_bullets = [
        "Designed and developed Subscription module with JWT-based authentication, API Gateway integration, and tiered access control (Spring Cloud Gateway, Spring Security).",
        "Developed backend services for Book Distributor platform using Spring Boot, Spring Data, MySQL, AWS, Hibernate with comprehensive REST APIs.",
    ]
    for b in oz_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    # Intern - DirectFN
    p = doc.add_paragraph()
    p.add_run("Software Engineer Intern").bold = True
    p.add_run(" | July 2019 – July 2020")
    p = doc.add_paragraph()
    p.add_run("DirectFN,").bold = True
    p.add_run(" Sri Lanka")
    intern_bullets = [
        "Developed POC backend system with secure, fault-tolerant architecture implementing RESTful services and event-driven design for financial data processing.",
        "Implemented encryption-based security enhancements for HSBC peer-to-peer communications using TLS/mTLS.",
    ]
    for b in intern_bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()

    add_heading(doc, "EDUCATION", level=1)
    p = doc.add_paragraph()
    p.add_run("University of Westminster, UK").bold = True
    doc.add_paragraph("BEng (Hons) Software Engineering – First Class")
    doc.add_paragraph()

    add_heading(doc, "PUBLICATIONS", level=1)
    doc.add_paragraph("An Analysis Of Face Recognition Under Face Mask Occlusions – MLDS Conference 2021, Zurich, Switzerland")

    out_path = "resume_optimized_senior_full_stack.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()
